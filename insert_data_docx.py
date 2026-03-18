from __future__ import annotations
import os
import sys
from pathlib import Path
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK


# ==========================
# 🔧 CONFIGURA TUS RUTAS AQUÍ
# ==========================

POPPLER_PATH = r"C:\Users\raquell.martinez\poppler\Release-25.12.0-0\poppler-25.12.0\Library\bin"

# ==========================


def validar_poppler():
    if not os.path.isdir(POPPLER_PATH):
        raise FileNotFoundError(
            f"No se encontró Poppler en:\n{POPPLER_PATH}"
        )


def insertar_primeras_paginas(docx_path: str, pdf_files: list[str]) -> None:
    validar_poppler()

    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"No existe el archivo Word: {docx_path}")

    doc = Document(docx_path)

    for pdf in pdf_files:
        try:
            if not os.path.isfile(pdf):
                print(f"⚠ No existe el PDF: {pdf}")
                continue

            print(f"Procesando: {pdf}")

            # Convertir SOLO la primera página
            pages = convert_from_path(
                pdf,
                first_page=1,
                last_page=1,
                poppler_path=POPPLER_PATH
            )

            if not pages:
                print(f"⚠ No se pudo convertir: {pdf}")
                continue

            temp_image = f"{Path(pdf).stem}_page1_temp.png"
            pages[0].save(temp_image, "PNG")

            # Salto de página
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

            # Insertar imagen
            doc.add_picture(temp_image, width=Inches(6))

            # Eliminar imagen temporal
            os.remove(temp_image)

            print(f"✔ Insertado correctamente: {pdf}")

        except Exception as e:
            print(f"❌ Error procesando {pdf}: {e}")

    # Guardar cambios en el mismo archivo
    doc.save(docx_path)
    print("\nDocumento actualizado correctamente.")


# ==========================
# 📌 CONFIGURA TUS ARCHIVOS
# ==========================

if __name__ == "__main__":

    docx_path = r"cv/final_cv/David Barrera Munoz - CV_generado.docx"

    pdf_files = [
        r"cv/templates_input/testers/Cedula_Prof-David Barrera Muñoz.pdf",
        r"cv/templates_input/testers/Título-DavidBarrera.pdf"
    ]

    try:
        insertar_primeras_paginas(docx_path, pdf_files)
    except Exception as e:
        print(f"Error general: {e}")
        sys.exit(1)